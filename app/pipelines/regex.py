import json
import re
from typing import Dict, Any, Optional
from datetime import datetime
import docx
import pdfplumber

class FinancialDocumentParser:
    def __init__(self):
        # Regular expressions to extract key financial entities
        self.entity_patterns = {
            'party_a': r'^Party A\s*(?::|\||\t)\s*(.+)$',
            'party_b': r'^Party B\s*(?::|\||\t)\s*(.+)$',
            'trade_date': r'^Trade Date\s*(?::|\||\t)\s*(\d{1,2}\s+\w+\s+\d{4})$',
            'trade_time': r'^Trade Time\s*(?::|\||\t)\s*(\d{2}:\d{2}:\d{2})$',
            'initial_valuation_date': r'^Initial Valuation Date\s*(?::|\||\t)\s*(\d{1,2}\s+\w+\s+\d{4})$',
            'effective_date': r'^Effective Date\s*(?::|\||\t)\s*(\d{1,2}\s+\w+\s+\d{4})$',
            'valuation_date': r'^Valuation Date\s*(?::|\||\t)\s*(\d{1,2}\s+\w+\s+\d{4})$',
            'termination_date': r'^(?:Termination Date|Maturity)\s*(?::|\||\t)\s*(\d{1,2}\s+\w+\s+\d{4})$',
            'notional_amount': r'^Notional Amount\s*\(N\)\s*(?::|\||\t)\s*([A-Z]{3}\s+[\d,.\s]+(?:million|thousand|bn|mm|m|k|b)?(?:\s+\w+)?)$',
            'upfront_payment': r'^Upfront Payment\s*(?::|\||\t)\s*(.+)$',
            'underlying': r'^Underlying\s*(?::|\||\t)\s*([^(]+)\(ISIN\s+([A-Z0-9]+),\s*Reuters:\s*([A-Z0-9.]+)\)\s*$',
            'exchange': r'^Exchange\s*(?::|\||\t)\s*([^\n]+)$',
            'coupon': r'^Coupon\s*\(C\)\s*(?::|\||\t)\s*([^\n]+)$',
            'barrier': r'^Barrier\s*\(B\)\s*(?::|\||\t)\s*([^\n]+)$',
            'interest_payments': r'^Interest Payments(?:\s*(?::|\||\t)\s*([^\n]+))?$',
            'initial_price': r'^Initial Price\s*\(Shareini\)\s*(?::|\||\t)\s*([^\n]+)$',
            'final_price': r'^Sharefinal\s*(?::|\||\t)\s*([^\n]+)$',
            'business_day': r'^Business Day\s*(?::|\||\t)\s*([^\n]+)$',
            'future_price_valuation': r'^Future Price Valuation\s*(?::|\||\t)\s*([^\n]+)$',
            'calculation_agent': r'^Calculation Agent\s*(?::|\||\t)\s*([^\n]+)$',
            'isda_doc': r'^ISDA Documentation\s*(?::|\||\t)\s*([^\n]+)$',
        }

    def extract_from_docx(self, file_path: str) -> Dict[str, Any]:

        try:
            doc = docx.Document(file_path)

            # Collect all paragraphs
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

            # Collect all table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text_parts.append('\t'.join(cells))

            full_text = self.preprocess_text('\n'.join(text_parts))
            return self.extract_and_structure_entities(full_text)

        except Exception as e:
            return {'error': f'Failed to process DOCX: {str(e)}'}

    def extract_from_text(self, text: str) -> Dict[str, Any]:

        return self.extract_and_structure_entities(self.preprocess_text(text))

    def extract_and_structure_entities(self, text: str) -> Dict[str, Any]:
        # First, extract raw entities using regex patterns
        raw = {}
        for key, pattern in self.entity_patterns.items():
            m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if not m:
                continue

            if m.lastindex and m.lastindex > 1:
                raw[key] = [self.clean_value(g) for g in m.groups()]
            else:
                raw[key] = self.clean_value(m.group(1)) if m.group(1) else None

        # Handle special case for interest payments
        if raw.get('interest_payments') in (None, "Interest Payments"):
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if re.match(r'^Interest Payments', line, re.IGNORECASE):
                    if i + 1 < len(lines):
                        raw['interest_payments'] = self.clean_value(lines[i + 1])
                    break

        structured = self.structure_entities(raw)

        return {
            'document_type': 'structured_financial_document',
            'extraction_timestamp': datetime.now().isoformat(),
            'entities': structured,
            'confidence_score': self.calculate_confidence(structured)
        }

    def structure_entities(self, r: Dict[str, Any]) -> Dict[str, Any]:
        # Map raw fields to structured output
        out = {
            'PartyA': r.get('party_a'),
            'PartyB': r.get('party_b'),
            'TradeDate': self.std_date(r.get('trade_date')),
            'TradeTime': r.get('trade_time'),
            'InitialValuationDate': self.std_date(r.get('initial_valuation_date')),
            'EffectiveDate': self.std_date(r.get('effective_date')),
            'ValuationDate': self.std_date(r.get('valuation_date')),
            'TerminationDate': self.std_date(r.get('termination_date')),
            'UpfrontPayment': r.get('upfront_payment'),
            'Coupon': r.get('coupon'),
            'InterestPayments': r.get('interest_payments'),
            'BusinessDay': r.get('business_day'),
            'FuturePriceValuation': r.get('future_price_valuation'),
            'ISDADocumentation': r.get('isda_doc')
        }

        # Parse notional amount
        if 'notional_amount' in r:
            notional = self.parse_notional(r['notional_amount'])
            if notional:
                out['Notional'] = notional

        # Parse underlying
        if 'underlying' in r and isinstance(r['underlying'], list) and len(r['underlying']) == 3:
            name, isin, ticker = r['underlying']
            out['Underlying'] = {
                'name': name,
                'isin': isin,
                'ticker': ticker,
                'exchange': r.get('exchange')
            }

        # Parse barrier
        if 'barrier' in r:
            val, ref = self.parse_barrier(r['barrier'])
            if val is not None:
                out['Barrier'] = {'value': val, 'unit': '%'}
                if ref:
                    out['Barrier']['reference'] = ref

        # Equity payments
        equity = {}
        if r.get('initial_price'):
            equity['InitialPrice'] = r.get('initial_price')
        if r.get('final_price'):
            equity['FinalPrice'] = r.get('final_price')
        if equity:
            out['EquityPayments'] = equity

        # Calculation agent(s)
        if r.get('calculation_agent'):
            agents = [a.strip() for a in re.split(r'\s*and\s*', r['calculation_agent'], flags=re.IGNORECASE)]
            out['CalculationAgent'] = agents if len(agents) > 1 else agents[0]

        # Remove None values
        return {k: v for k, v in out.items() if v is not None}

    def preprocess_text(self, s: str) -> str:
        # Normalize line endings and whitespace
        s = s.replace('\r\n', '\n').replace('\r', '\n')
        s = re.sub(r'[ \t]+', lambda m: '\t' if '\t' in m.group(0) else ' ', s)
        return '\n'.join(line.strip() for line in s.split('\n') if line.strip())

    def clean_value(self, v: str) -> str:
        v = v.strip()
        v = re.sub(r'^\|+\s*', '', v)
        return re.sub(r'\s+', ' ', v).strip()

    def std_date(self, date_str: Optional[str]) -> Optional[str]:
        if not date_str:
            return None
        try:
            return datetime.strptime(date_str, "%d %B %Y").strftime("%Y-%m-%d")
        except ValueError:
            return date_str

    def parse_notional(self, s: str) -> Optional[Dict[str, Any]]:
        m = re.search(
            r'^(?P<ccy>[A-Z]{3})\s+(?P<num>[\d.,]+|\d+(?:\.\d+)?)\s*(?P<scale>million|thousand|bn|mm|m|k|b)?(?:\s+(?P<unit>\w+))?$',
            s.strip(), re.IGNORECASE
        )
        if not m:
            return None

        ccy, num_raw, scale, unit = m.group('ccy').upper(), m.group('num').replace(',', ''), (m.group('scale') or '').lower(), m.group('unit')
        try:
            num = float(num_raw)
        except ValueError:
            return None

        mult = {
            'thousand': 1e3, 'k': 1e3,
            'million': 1e6, 'm': 1e6, 'mm': 1e6,
            'bn': 1e9, 'b': 1e9, 'billion': 1e9
        }.get(scale, 1)

        return {'amount': num * mult, 'currency': ccy, **({'unit': unit} if unit else {})}

    def parse_barrier(self, s: str):
        m = re.search(r'(\d+(?:\.\d+)?)\s*%', s.strip())
        val = float(m.group(1)) if m else None

        ref = None
        m2 = re.search(r'of\s+([A-Za-z0-9_]+)', s, re.IGNORECASE)
        if m2:
            ref = m2.group(1)

        return val, ref

    def calculate_confidence(self, entities: Dict) -> float:
        key_fields = ['PartyA', 'PartyB', 'Notional', 'Underlying', 'TerminationDate', 'TradeDate']
        found = sum(1 for f in key_fields if f in entities)
        return round(found / len(key_fields), 2)

    def export_results(self, results: Dict, output_path: str, format='json'):
        if format == 'json':
            with open(f"{output_path}.json", 'w') as f:
                json.dump(results, f, indent=2)
\
def run_regex(file):
    """
    file: can be UploadFile (from FastAPI) or plain text string
    """
    parser = FinancialDocumentParser()

    # If it's a FastAPI UploadFile
    try:
        filename = file.filename
        content_type = file.content_type
    except AttributeError:
        # It's already a text string
        return parser.extract_from_text(file)

    # PDF file
    if filename.lower().endswith(".pdf") or content_type == "application/pdf":
        with pdfplumber.open(file.file) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return parser.extract_from_text(text)

    # DOCX file
    elif filename.lower().endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parser.extract_from_docx(file.file)

    # TXT file
    elif filename.lower().endswith(".txt") or content_type == "text/plain":
        content = file.file.read()
        text_content = content.decode("utf-8", errors="ignore")
        return parser.extract_from_text(text_content)

    else:
        raise ValueError("Unsupported file type")
